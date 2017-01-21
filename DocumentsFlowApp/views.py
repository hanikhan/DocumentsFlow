import io
import os
import subprocess
import shutil
import uuid

from django.contrib.auth import authenticate, login, logout
from django.core.files.base import ContentFile
from django.core.files.storage import default_storage
from django.shortcuts import render, redirect, render_to_response
from django.http import HttpResponseRedirect
from django.utils import timezone
from django.utils.encoding import smart_str
from django.views.static import serve

from DocumentsFlowApp.models import Document, MyUser, Task, Template
from utils import get_project_path_forward_slash, get_project_path
from .forms import UploadFileForm, UploadNewVersionForm
from .forms import CreateFileForm
from DocumentsFlow.settings import TEMPLATES_PATH

# Create your views here.

from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.template.context_processors import csrf
from django.views.decorators.csrf import csrf_protect
import datetime
from django.http import JsonResponse
from DocumentsFlowApp.models import Document
from DocumentsFlowApp.models import Log
from DocumentsFlowApp.models import *
from decimal import Decimal
from django.core.mail import send_mail

from docx import Document as WordDocument


def index(request):
    return render(request, "login.djt")


@csrf_protect
def homepage(request):
    print("******* " + str(request))
    c = {}
    c.update(csrf(request))
    if "user" not in str(request):
        try:
            user = authenticate(username=request.POST.get('username'), password=request.POST.get('password'))
            if user is not None:
                login(request, user)
                return render(request, "homepage2.html", c)
        except Exception:
            return render(request, "login.djt")
    else:
        print("here")
        return render(request, "homepage2.html", c)


@login_required
@csrf_protect
def zona_de_lucru(request):
    c = {}
    c.update(csrf(request))

    user_docs = []
    docs = Document.objects.all()
    print(request.user)
    for doc in docs:
        if doc.get_owner().username == request.user.username and doc.get_task().id == 110:
            wasDeleted = deleteDocumentAfter30(request, doc)
            if wasDeleted == True:
                continue
            if doc.get_task().id == 110:
                user_docs.append(doc)
    c["docs"] = user_docs
    return render(request, "zona_de_lucru.html", c)


def deleteDocumentAfter30(request, document):
    if (datetime.datetime.now().date() - document.get_date()).days > 30:
        log = Log()
        log.set_date(datetime.datetime.now().date())
        log.set_action("PROGRAMAT PENTRU STERGERE")
        log.set_document_name(document.get_name())
        log.set_document_path(document.get_path())
        log.set_user(document.get_owner())
        log.save()

        send_mail("Delete", "The file " + document.get_name() + " with version " + str(
            document.get_version()) + " will be deleted in 30 days", "websmarts2017@gmail.com", [request.user.email],
                  fail_silently=False)
        return False
    elif (datetime.datetime.now().date() - document.get_date()).days > 60:
        log = Log()
        log.set_date(datetime.datetime.now().date())
        log.set_action("STERGERE DOCUMENT")
        log.set_document_name(document.get_name())
        log.set_document_path(document.get_path())
        log.set_user(document.get_owner())
        log.save()

        send_mail("Delete",
                  "The file " + document.get_name() + " with version " + str(document.get_version()) + " was deleted",
                  "websmarts2017@gmail.com", [request.user.email], fail_silently=False)
        document.delete()
        default_storage.delete(document.get_path())
        return True
    return False


@csrf_protect
def change_document_status_to_final(request):
    document_path = request.GET.get("path")
    document = Document.objects.filter(path=document_path).first()
    document.set_status("FINAL")
    document.set_version(1.0)
    document.save()

    #default_storage.delete(document.get_path())
    new_path=get_project_path_forward_slash() + "resources/"+str(document.get_version())+'^'+(document_path.split('^',1))[1]
    print(new_path)

    os.rename(document.get_path(), new_path)
    #default_storage.save(new_path,file)

    document.set_path(new_path)
    document.save()


    log = Log()
    log.set_date(datetime.datetime.now().date())
    log.set_action("MAKE DOCUMENT FINAL")
    log.set_document_name(document.get_name())
    log.set_document_path(document.get_path())
    log.set_user(document.get_owner())
    log.save()

    c = {}
    c.update(csrf(request))
    return render(request, "homepage2.html", c)


@csrf_protect
def change_document_status_to_draft(request):
    document_path = request.GET.get("path")
    document = Document.objects.filter(path=document_path).first()
    document.set_status("DRAFT")
    document.set_version(0.1)
    document.save()

    # default_storage.delete(document.get_path())
    new_path = get_project_path_forward_slash() + "resources/" + str(document.get_version()) + '^' + \
               (document_path.split('^', 1))[1]
    print(new_path)

    os.rename(document.get_path(), new_path)
    # default_storage.save(new_path,file)

    document.set_path(new_path)
    document.save()


    log = Log()
    log.set_date(datetime.datetime.now().date())
    log.set_action("MAKE DOCUMENT DRAFT")
    log.set_document_name(document.get_name())
    log.set_document_path(document.get_path())
    log.set_user(document.get_owner())
    log.save()

    c = {}
    c.update(csrf(request))
    return render(request, "homepage2.html", c)


@csrf_protect
def delete_draft(request):
    document_path = request.POST.get("document_path")
    document_path = document_path.replace("\\", "\\\\")
    document = Document.objects.filter(path=document_path).first()
    log = Log()
    log.set_date(datetime.datetime.now().date())
    log.set_action("STERGERE DRAFT")
    log.set_document_name(document.get_name())
    log.set_document_path(document.get_path())
    log.set_user(document.get_owner())
    log.save()
    document.delete()
    default_storage.delete(document.get_path())

    c = {}
    c.update(csrf(request))
    return render(request, "homepage2.html", c)


def logout_user(request):
    logout(request)
    return redirect("/")


def add_uploaded_file(request, file, abstract, keywords):

      # for item in Document.objects.all():
      #         item.delete()


    ok = False
    print(get_project_path_forward_slash())
    for item in Document.objects.all():
        path = item.get_path().split("^", 1)
        if item.get_owner().username == request.user.username \
                and get_project_path_forward_slash() + "resources/" + path[1] == \
                    get_project_path_forward_slash() + "resources" + "/" + request.user.username + file.name:
            if item.get_status() == "DRAFT":
                doc = Document.objects.filter(name=file.name).last()
                newDocument = Document()
                newDocument.set_name(file.name)
                newDocument.set_status("DRAFT")
                newDocument.set_version(doc.get_version() + Decimal('0.1'))
                newDocument.set_date(datetime.datetime.now())
                newDocument.set_owner(request.user)
                path = default_storage.save(get_project_path_forward_slash() + "resources/" +
                                            str(newDocument.get_version()) + "^" + path[1], file)
                print("*******" + str(path))
                newDocument.set_path(path)
                ts = Task.objects.all()
                for t in ts:
                    if t.id == 110:
                        newDocument.set_task(t)
                        break

                newDocument.set_abstract(abstract)
                newDocument.set_keywords(keywords)

                log = Log()
                log.set_date(datetime.datetime.now().date())
                log.set_action("UPLOAD DOCUMENT")
                log.set_document_name(newDocument.get_name())
                log.set_document_path(newDocument.get_path())
                log.set_user(newDocument.get_owner())
                log.save()

                newDocument.save()
                ok = True
                break

            if item.get_status() == "FINAL":
                item.set_version(item.get_version() + Decimal('1.0'))
                default_storage.delete(item.get_path())
                path = default_storage.save(item.get_path(), file)
                item.set_path(path)
                item.set_date(datetime.datetime.now())
                ts = Task.objects.all()
                for t in ts:
                    if t.id == 110:
                        item.set_task(t)
                        break

                item.set_abstract(abstract)
                item.set_keywords(keywords)

                log = Log()
                log.set_date(datetime.datetime.now().date())
                log.set_action("UPLOAD DOCUMENT")
                log.set_document_name(item.get_name())
                log.set_document_path(item.get_path())
                log.set_user(item.get_owner())
                log.save()

                item.save()
                ok = True
            break

    if ok == False:
        d = Document()
        d.set_name(file.name)
        d.set_version(0.1)
        d.set_owner(request.user)
        d.set_status("DRAFT")
        d.set_date(datetime.datetime.now())
        d.set_keywords(keywords)
        d.set_abstract(abstract)
        ts = Task.objects.all()
        for t in ts:
            if t.id == 110:
                d.set_task(t)
                break
        path = default_storage.save(
            get_project_path() + "resources" + "\\" + str(d.get_version()) + "^" + request.user.username + file.name,
            file)
        d.set_path(path)

        log = Log()
        log.set_date(datetime.datetime.now().date())
        log.set_action("UPLOAD DOCUMENT")
        log.set_document_name(d.get_name())
        log.set_document_path(d.get_path())
        log.set_user(d.get_owner())
        log.save()

        d.save()


@csrf_protect
def upload_file(request):
    c = {}
    c.update(csrf(request))

    if request.method == 'POST':
        form = UploadFileForm(request.POST, request.FILES)
        print(request.POST.get('abstract'))
        if form.is_valid():
            add_uploaded_file(request, request.FILES['file'],request.POST.get('abstract'),request.POST.get('keywords'))
            return render(request, "homepage2.html", c)
    else:
        form = UploadFileForm()
    return render(request, "uploadFile.html", {'form': form})


def approve_task(document_id):
    document = Document.objects.filter(document_id=document_id).first()
    task = document.get_task()
    next_step = task.get_step()+1

    assignment_id = task.get_assigment().id

    process_id = task.get_process().id

    create_task_for_process(assignment_id,process_id,next_step)

def accept_task(request):
    task_id = int(request.GET.get("task_id"))
    task = Task.objects.filter(id=int(task_id)).first()
    task.set_status("ACCEPTED")
    flux = task.get_process().get_flux()
    counter = 0
    for assig in Assigment.objects.all():
        if assig.get_flux() == flux:
            counter +=1

    if task.get_step() != counter:
        start_task_for_step(task.get_process().id,task.get_step()+1)
    else:
        process = task.get_process()
        process.set_status("APPROVED")
        process.save()

        for doc in Document.objects.all():
            if doc.get_owner() == task.get_process().get_starter():
                doc.set_status("BLOCAT")
                doc.save()

    return zona_taskuri(request)

def respinge_task(request):

    task_id = int(request.GET.get("task_id"))
    task = Task.objects.filter(id=task_id).first()
    task.set_status("REJECTED")
    process = task.get_process()
    process.set_status("REJECTED")
    process.save()

    for doc in Document.objects.all():
        if doc.get_owner() == task.get_process().get_starter():
            doc.set_status("BLOCAT")
            doc.save()

    for taskk in Task.objects.all():
        if taskk.get_process() == process and task.get_step() < taskk.get_step():
            taskk.delete()

    return zona_taskuri(request)


def start_task_for_step(process_id,step):
     process = Process.objects.filter(id = process_id).first()
     flux = Flux.objects.filter(id = process.get_flux().id).first()
     deadline_days = -1;
     for assignment in Assigment.objects.all():
         if assignment.get_flux() == flux and assignment.get_step() == step:
             deadline_days = assignment.get_days()

     new_task = Task()
     for task in Task.objects.all():
         if task.get_step() == step:
             task.set_status("PENDING")
             task.set_deadline(datetime.datetime.now().date()+datetime.timedelta(days=deadline_days))
             new_task = task

     new_task.save()

     for document in Document.objects.all():
         if document.get_task().get_step() == step-1 and document.get_task().get_process() == process:
             document.set_task(new_task)
             document.save()


def add_document_to_process(request):
    document_id = request.GET.get('document_id')
    process_id = request.GET.get('procces_id')
    process1=Process.objects.filter(id=int(process_id)).first()
    document = Document.objects.filter(id = int(document_id)).first()
    tasks = Task.objects.filter(process=process1)

    for task in tasks:
        if task.get_step() == 1:
            document.set_task(task)

    document.save()

    c = {}
    c.update(csrf(request))

    user_docs = []
    docs = Document.objects.all()
    for doc in docs:
        if doc.get_owner().username == request.user.username:
            if doc.get_status() == "FINAL":
                user_docs.append(doc)
    c["docs"] = user_docs
    c["process_id"] = process1.id

    return render(request, "process.html", c)

def start_process(request):
    process_id = int(request.GET.get('process_id'))
    process = Process.objects.filter(id = process_id).first()
    process.set_status("activ");
    process.save()
    start_task_for_step(process_id,1)


    c = {}
    c.update(csrf(request))

    user_docs = []
    docs = Document.objects.all()
    for doc in docs:
        if doc.get_owner().username == request.user.username:
            if doc.get_status() == "FINAL":
                user_docs.append(doc)
    c["docs"] = user_docs
    c["process_id"] = process_id

    return render(request, "process.html", c)

def create_task_for_process(assignment_id, process_id, step):
    assignment = Assigment.objects.filter(id = assignment_id).first();
    process = Process.objects.filter(id = process_id).first();
    task = Task()
    task.set_assigment(assignment)
    task.set_process(process)
    task.set_status("NOT_STARTED")
    days = assignment.get_days()
    task.set_deadline(datetime.datetime.now().date())
    task.set_step(step)

    task.save()

    return task.id


def create_procces(request,flux_id):

    p = Process()
    p.set_starter(request.user)
    flux = Flux.objects.filter(id=int(flux_id)).first()
    p.set_flux(flux)
    p.save()

    for assignment in Assigment.objects.all():
        if assignment.get_flux().id == int(flux_id):
            create_task_for_process(assignment.id, p.id, assignment.get_step())

    return p.id


@csrf_protect
def create_file(request):
    c = {}
    c.update(csrf(request))

    if request.method == 'POST':
        form = CreateFileForm(request.POST)
        if form.is_valid():
            docTitle = request.POST['doctitle'].replace(" ", "")
            document = WordDocument()
            docx_title = docTitle + ".docx"
            document.add_heading('Cerere', 1)
            document.add_paragraph('Pentru efectuarea deplasarilor in strainatate(eliberare Dispozitia Rectorului)')
            document.add_paragraph(
                'Subsemnatul(a) ' + request.POST['name'] + ' avand titulatura/functia de: ' + request.POST['function']
                + ' in cadrul Facultatii/Departamentului(pentru personal administrativ) de' + request.POST['department']
                + ', deplasarea la: (se precizeaza localitatea*/tara) '+ request.POST['to']
            )

            document.add_paragraph('Ruta: ' + request.POST['route'])
            document.add_paragraph('Perioada cand are loc actiunea**: ' + request.POST['period'])
            document.add_paragraph('Perioada de deplasare: ' + request.POST['periodOfDeparture'])
            document.add_paragraph('Cu mijloc de transport: ' + request.POST['vehicle'])
            document.add_paragraph(
                'Numar de telefon ' + request.POST['phoneNumber']
                + ' adresa de email ' + request.POST['email'] + '.'
            )

            document.add_paragraph('Scopul deplasarii: ' + request.POST['scopeOfDeparture'])
            document.add_paragraph('Cheltuielile aferente mobilitatii sunt suportate de: ' + request.POST['expenses'])
            document.add_paragraph('Cheltuieli de deplasare solicitate:')
            document.add_paragraph('1.Diurna/Subzistenta')
            document.add_paragraph(
                'Cuantul ' + request.POST['scopeOfDeparture'] + ' /pe zi, nr zile ' + request.POST['nrOfDays']
                + ' ,total ' + request.POST['total']
            )

            document.add_paragraph('2.Cazare')
            document.add_paragraph(
                'Cuantul ' + request.POST['accomodationCuantum'] + ' /pe zi, nr zile ' + request.POST['accomodationNrOfDays']
                + ' ,total ' + request.POST['accomodationTotal']
            )

            temp_path = get_project_path() + "resources\\\\0.1^" + request.user.username + docx_title
            path = get_project_path() + "resources\\\\0.1^" + request.user.username + docx_title
            print("**** " + str(path))
            document.save(temp_path)
            document_model = Document()
            document_model.set_date(timezone.now())
            document_model.set_name(docx_title)
            document_model.set_owner(request.user)
            document_model.set_path(path)
            document_model.set_status("DRAFT")
            for template in Template.objects.all():
                if template.id == 1:
                    document_model.set_template(template)
                    document_model.set_template_values(template.get_keys())
                    break
            for task in Task.objects.all():
                if task.id == 110:
                    document_model.set_task(task)
                    break
            document_model.set_type("")
            document_model.set_version(0.1)

            log = Log()
            log.set_date(datetime.datetime.now().date())
            log.set_action("CREATE DOCUMENT")
            log.set_document_name(document_model.get_name())
            log.set_document_path(document_model.get_path())
            log.set_user(document_model.get_owner())
            log.save()

            document_model.save()

            return render(request, "homepage2.html", c)
            # length = f.tell()
            # f.seek(0)
            # response = HttpResponse(
            #     f.getvalue(),
            #     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            # )
            # response['Content-Disposition'] = 'attachment; filename=' + docx_title
            # response['Content-Length'] = length
            # return response
    else:
        form = CreateFileForm()
    return render(request, "createFile.html", {'form': form})


def download_file(request):
    file_path = request.GET.get("path")
    print(">>>>>> " + str(file_path))
    fsock = open(file_path, "rb")
    doc_name = request.GET.get("doc_name")
    response = HttpResponse(fsock, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=%s' % smart_str(doc_name)
    response['X-Sendfile'] = smart_str(file_path)
    return response


def edit_metadata(request):
    return render(request, "editMetadata.html")


def processes(request):
    c = {}
    c.update(csrf(request))

    fluxes = Flux.objects.all()
    c["fluxes"] = fluxes
    return render(request, "processes.html", c)


def process(request):
    c = {}
    c.update(csrf(request))

    user_docs = []
    docs = Document.objects.all()
    for doc in docs:
        if doc.get_owner().username == request.user.username:
            if doc.get_status()=="FINAL" and doc.get_task().id==110:
                user_docs.append(doc)
    c["docs"] = user_docs
    flux_id = request.GET.get('flux_id')
    procces_id =create_procces(request,flux_id)
    c["process_id"] = procces_id
    return render(request, "process.html", c)



def zona_taskuri_initiate(request):
    c = {}
    c.update(csrf(request))

    user_docs = []
    docs = Document.objects.all()
    print(request.user)
    for doc in docs:
        if doc.get_owner().username == request.user.username and doc.get_task().id != 110 and doc.get_task().get_process().get_status() == "activ":
            user_docs.append(doc)
    c["docs"] = user_docs
    return render(request, "zona_taskuri_initiate.html", c)


def zona_taskuri(request):
    c = {}
    c.update(csrf(request))

    user_docs = []
    docs = Document.objects.all()

    user_group = request.user.get_group().id

    user_assigments = []
    for assigment in Assigment.objects.all():
        if assigment.get_user() != 2 and assigment.get_user().get_full_name() != "du" and assigment.get_user().get_full_name() == request.user.username:
            user_assigments.append(assigment)
        elif assigment.get_board() == user_group:
            user_assigments.append(assigment)

    user_tasks = []
    for assig in user_assigments:
        for task in Task.objects.all():
            if task.get_assigment() == assig and task.get_status() == "PENDING" and task.get_process().get_status() == "activ":
                user_tasks.append(task)


    for doc in Document.objects.all():
        for task in user_tasks:
            if doc.get_task() == task:
                user_docs.append(doc)

    c["docs"] = user_docs
    return render(request, "zona_taskuri.html", c)


def logs(request):
    c = {}
    c.update(csrf(request))

    logs = Log.objects.all()
    logs_list = []
    for log in logs:
        logs_list.append({"user": log.get_user().username,
                          "action": log.get_action(),
                          "docName": log.get_document_name(),
                          "docPath": log.get_document_path(),
                          "date": log.get_date()})
    c["logs"] = logs_list
    return render(request, "logs.html", c)


def zona_taskuri_terminate(request):
    c = {}
    c.update(csrf(request))

    user_docs = []
    docs = Document.objects.all()
    print(request.user)
    for doc in docs:
        if doc.get_owner().username == request.user.username and doc.get_status()=="BLOCAT":
            user_docs.append(doc)

    c["docs"] = user_docs
    return render(request, "zona_taskuri_terminate.html", c)


def cancel_process(request):
    process_id = int(request.GET.get("process_id"))
    process = Process.objects.filter(id=process_id).first()
    process.delete()
    return processes(request)


def uploadNewVersion(request):
    c = {}
    c.update(csrf(request))

    if request.method == 'POST':
        form = UploadNewVersionForm(request.POST, request.FILES)
        if form.is_valid():
            upload_final_revizuit(request, request.FILES['file'])
            return render(request, "homepage2.html", c)
    else:
        form = UploadNewVersionForm()
    return render(request, "uploadNewVersion.html", {'form': form})




def upload_final_revizuit(request, file):

    for item in Document.objects.all():
        path = item.get_path().split("^", 1)

        if item.get_owner().username != request.user.username \
                and get_project_path_forward_slash() + "resources/" + path[1] == \
                    get_project_path_forward_slash() + "resources" + "/" + item.get_owner().username + file.name:
            if item.get_status()=="FINAL":
                item.set_status("FINAL REVIZUIT")
                item.set_version(item.get_version() + Decimal('0.1'))
                default_storage.delete(item.get_path())
                print(get_project_path_forward_slash() + "resources/" +str(item.get_version()) + "^" + path[1])
                item.set_path(get_project_path_forward_slash() + "resources/" +str(item.get_version()) + "^" + path[1])
                item.save()
                path = default_storage.save(item.get_path(), file)
                break
            else:
                item.set_version(item.get_version() + Decimal('0.1'))
                default_storage.delete(item.get_path())
                print(get_project_path_forward_slash() + "resources/" + str(item.get_version()) + "^" + path[1])
                item.set_path(get_project_path_forward_slash() + "resources/" +str(item.get_version()) + "^" + path[1])
                path = default_storage.save(item.get_path(), file)
                item.save()

def pdf_view(request):
    document_id = int(request.GET.get('document_id'))
    document = Document.objects.filter(id=document_id).first()
    document_path = document.get_path()


    with open(document_path, 'rb') as pdf:
        response = HttpResponse(pdf.read(), content_type='application/pdf')
        response['Content-Disposition'] = 'inline;filename=somefile.pdf'
        return response
    extension.closed


def create_flux(request):
    print("here")
    if request.method == 'GET':
        selected_templates = request.GET.getlist("templates[]")
        selected_group = request.GET.get("group", None)
        get_user = request.GET.get("get_user", None)
        current_step = request.GET.get("current_step", 1)
        gdone = request.GET.get("gdone", None)
        gcontinue = request.GET.get("gcontinue", None)
        udone = request.GET.get("udone", None)
        ucontinue = request.GET.get("ucontinue", None)
        if len(selected_templates) > 0:
            print("in select group")
            flux_name = request.GET.get("fluxName")
            flux = Flux()
            docs = ""
            for doc in selected_templates:
                docs += doc + ","
                flux.set_name(flux_name)
            flux.set_documents(docs)
            # flux.save()

            groups = []
            for group in Group.objects.all():
                if not  group.get_name() == "AdminGroup":
                    groups.append({"name": group.get_name(), "id":group.id})

            data = {"flux_id": 2,
                    "selected_docs": docs,
                    "groups": groups,
                    "current_step": 1}
            return render(request, 'createFlux.html', data)
        elif get_user is not None:
            print("in get user")
            flux_id = request.GET.get("flux_id")
            users = MyUser.objects.filter(group=int(selected_group))
            formatted_users = []
            for user in users:
                formatted_users.append(user.username)
            response = {"users": formatted_users,
                        "current_step": current_step,
                        "flux_id": flux_id}
            return render(request, 'createFlux.html', response)
        elif ucontinue:
            print("in user continue")
            flux_id = request.GET.get("flux_id")
            days = request.GET.get("days")
            flux = Flux.objects.filter(id=flux_id)[0]
            user_id = request.GET.get("user", None)
            user = MyUser.objects.filter(username=user_id)[0]
            group_id = request.GET.get("group", None)
            assignment = Assigment()
            assignment.set_document_types(flux.get_documents())
            assignment.set_flux(flux)
            assignment.set_days(days)
            assignment.set_step(current_step)
            assignment.set_user(user)
            assignment.save()

            groups = []
            for group in Group.objects.all():
                groups.append({"name": group.get_name(), "id": group.id})
            current_step = int(current_step) + 1
            response = {"flux_id": flux_id,
                        "groups": groups,
                        "current_step": current_step}
            return render(request, 'createFlux.html', response)
        elif udone:
            print("in user done")
            flux_id = request.GET.get("flux_id")
            days = request.GET.get("days")
            flux = Flux.objects.filter(id=flux_id)[0]
            user_id = request.GET.get("user", None)
            user = MyUser.objects.filter(username=user_id)[0]
            group_id = request.GET.get("group", None)
            assignment = Assigment()
            assignment.set_document_types(flux.get_documents())
            assignment.set_flux(flux)
            assignment.set_days(days)
            assignment.set_step(current_step)
            assignment.set_user(user)
            assignment.save()
            return render(request, 'homepage2.html')
        elif gcontinue:
            print("in group continue")
            flux_id = request.GET.get("flux_id")
            days = request.GET.get("days")
            flux = Flux.objects.filter(id=flux_id)[0]
            group_id = request.GET.get("group", None)

            default_user = MyUser.objects.filter(username="du")[0]

            assignment = Assigment()
            assignment.set_document_types(flux.get_documents())
            assignment.set_flux(flux)
            assignment.set_days(days)
            assignment.set_step(current_step)
            assignment.set_user(default_user)
            assignment.set_board(int(group_id))
            assignment.save()

            groups = []
            for group in Group.objects.all():
                groups.append({"name": group.get_name(), "id": group.id})
            current_step = int(current_step) + 1
            response = {"flux_id": flux_id,
                        "groups": groups,
                        "current_step": current_step}
            return render(request, 'createFlux.html', response)
        elif gdone:
            print("in group done")
            flux_id = request.GET.get("flux_id")
            days = request.GET.get("days")
            flux = Flux.objects.filter(id=flux_id)[0]
            group_id = request.GET.get("group", None)

            default_user = MyUser.objects.filter(username="du")[0]

            assignment = Assigment()
            assignment.set_document_types(flux.get_documents())
            assignment.set_flux(flux)
            assignment.set_days(days)
            assignment.set_step(current_step)
            assignment.set_user(default_user)
            assignment.set_board(int(group_id))
            assignment.save()

            return render(request, 'homepage2.html')
        else:
            print("here2")
            template_names = []
            for template in Template.objects.all():
                template_names.append(template.get_name())
            response = {'templates': template_names}
            return render(request, 'createFlux.html', response)


def filter_log(request):
    column = request.GET.get("column")
    value = request.GET.get("value")

    logs_list = []
    if column == "Action":
        logs = Log.objects.all()
        for log in logs:
            if value.upper() in log.get_action().upper():
                logs_list.append({"user": log.get_user().username,
                                  "action": log.get_action(),
                                  "docName": log.get_document_name(),
                                  "date": log.get_date()})
                print(log.get_action())
    elif column == "Username":
        logs = Log.objects.all()
        for log in logs:
            if value.upper() in log.get_user().username.upper():
                logs_list.append({"user": log.get_user().username,
                                  "action": log.get_action(),
                                  "docName": log.get_document_name(),
                                  "date": log.get_date()})
                print(log.get_action())
    elif column == "Document Name":
        logs = Log.objects.all()
        for log in logs:
            if value.upper() in log.get_document_name().upper():
                logs_list.append({"user": log.get_user().username,
                                  "action": log.get_action(),
                                  "docName": log.get_document_name(),
                                  "date": log.get_date()})
                print(log.get_action())
    elif column == "Date":
        logs = Log.objects.all()
        for log in logs:
            if value.upper() in str(log.get_date()).upper():
                logs_list.append({"user": log.get_user().username,
                                  "action": log.get_action(),
                                  "docName": log.get_document_name(),
                                  "date": log.get_date()})
                print(log.get_action())
    else:
        logs = []

    json_logs = {"logs": logs_list}

    return render(request, 'logs.html', json_logs)
